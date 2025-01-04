import os
import numpy as np
import sounddevice as sd
import librosa
import tensorflow as tf
import docx
from sklearn.preprocessing import LabelEncoder

# Load the trained model and label encoder
model = tf.keras.models.load_model('D:\\Uni Work\\Python\\Speech Recognition\\my_model.keras')
label_classes = np.load('D:\\Uni Work\\Python\\Speech Recognition\\label_encoder_classes.npy', allow_pickle=True)
label_encoder = LabelEncoder()
label_encoder.classes_ = label_classes

# Define constants based on training
MAX_TIMESTEPS = 157  # Use the max_timesteps from training
FEATURES = 87        # Confirm if this matches the feature dimension used during training

def record_audio(duration=5, fs=16000):
    print("Recording...")
    audio_data = sd.rec(int(duration * fs), samplerate=fs, channels=1, dtype='float32')
    sd.wait()
    print("Recording complete.")
    return audio_data.flatten()

def preprocess_audio(audio_data):
    # Generate Mel spectrogram with the fixed number of mel bins (FEATURES)
    spectrogram = librosa.feature.melspectrogram(y=audio_data, sr=16000, n_mels=FEATURES)
    spectrogram = np.log(spectrogram + 1e-9)  # Apply log scaling

    # Truncate or pad the spectrogram to ensure consistent time steps
    if spectrogram.shape[1] > MAX_TIMESTEPS:
        spectrogram = spectrogram[:, :MAX_TIMESTEPS]  # Truncate time steps
    elif spectrogram.shape[1] < MAX_TIMESTEPS:
        padding = MAX_TIMESTEPS - spectrogram.shape[1]
        spectrogram = np.pad(spectrogram, ((0, 0), (0, padding)), mode='constant', constant_values=0)  # Pad if too short
    
    return spectrogram

def predict_audio(model, spectrogram):
    # Reshape spectrogram to match model's input shape (1, timesteps, features)
    spectrogram = spectrogram.reshape(1, MAX_TIMESTEPS, FEATURES)
    prediction = model.predict(spectrogram)
    predicted_class = np.argmax(prediction, axis=1)
    return label_encoder.inverse_transform(predicted_class)[0]

def main():
    doc = docx.Document()
    doc.add_heading('Real-time Voice Recognition', level=1)
    
    while True:
        audio_data = record_audio(duration=5)
        spectrogram = preprocess_audio(audio_data)
        predicted_label = predict_audio(model, spectrogram)
        
        doc.add_paragraph(predicted_label)
        print("Predicted Label:", predicted_label)
        doc.save('D:\\Uni Work\\Python\\Speech Recognition\\VoiceRecognitionOutput.docx')

if __name__ == "__main__":
    main()

